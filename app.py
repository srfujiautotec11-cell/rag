import streamlit as st
import os
import io
import glob
from docx import Document
from config import config
from services.document_processor import document_processor
from services.embeddings import embedding_service
from services.generator import generator_service
from services.vector_store import vector_store

st.set_page_config(
    page_title="Personal Knowledge Base",
    page_icon="📚",
    layout="wide"
)

st.title("📚 Personal Knowledge Base RAG System")

# --- Helper Functions for Checkbox Toggling ---
def toggle_all_docs(doc_ids):
    """Callback to select/deselect all documents in the Documents tab."""
    select_all_status = st.session_state.get('select_all_docs', False)
    for doc_id in doc_ids:
        st.session_state[f'doc_select_{doc_id}'] = select_all_status

def toggle_all_output_files(file_paths):
    """Callback to select/deselect all files in the Document Ingestion tab."""
    select_all_status = st.session_state.get('select_all_output', False)
    for file_path in file_paths:
        st.session_state[f'del_{file_path}'] = select_all_status

if not config.GEMINI_API_KEY:
    st.error("⚠️ GEMINI_API_KEY not found! Please set it in your environment.")
    st.info(
        "To run locally, create a file named `.env` in the root directory and add your key:\n\n"
        "```\n"
        "GEMINI_API_KEY=your_api_key_here\n"
        "```"
    )
    st.stop()

tab1, tab2, tab3 = st.tabs(["📄 Documents", "💬 Ask Questions", "📥 Document Ingestion"])

with tab1:
    st.header("Upload Documents")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_files = st.file_uploader(
            "Choose files",
            type=config.ALLOWED_FILE_TYPES,
            help="Upload PDF, text, markdown, or Word documents",
            accept_multiple_files=True
        )
    
    with col2:
        category = st.text_input(
            "Category",
            placeholder="e.g., Python, Family, Personal",
            help="Tag your document for organization",
            key="doc_category"
        )
    
    if st.button("Upload & Process", type="primary"):
        if not uploaded_files:
            st.error("Please select at least one file to upload.")
        else:
            total_files = len(uploaded_files)
            progress_bar = st.progress(0, text="Starting upload...")

            for i, uploaded_file in enumerate(uploaded_files):
                progress_text = f"Processing file {i + 1}/{total_files}: {uploaded_file.name}"
                progress_bar.progress((i) / total_files, text=progress_text)
                
                try:
                    file_bytes = uploaded_file.getvalue()
                    file_type = uploaded_file.name.split('.')[-1].lower()
                    file_hash = document_processor.compute_file_hash(file_bytes)
                    
                    file_stream = io.BytesIO(file_bytes)
                    text = document_processor.extract_text(file_stream, file_type)
                    chunks = document_processor.chunk_text(text)
                    
                    embeddings = embedding_service.generate_embeddings_batch(chunks)
                    
                    document_id = vector_store.store_document(
                        filename=uploaded_file.name,
                        file_type=file_type,
                        category=st.session_state.doc_category or "Uncategorized",
                        file_hash=file_hash
                    )
                    
                    vector_store.store_chunks(document_id, chunks, embeddings)
                    
                    st.success(f"✅ Successfully processed '{uploaded_file.name}'")
                    
                except Exception as e:
                    st.error(f"Error processing '{uploaded_file.name}': {str(e)}")

            progress_bar.progress(1.0, text="All files processed!")
            st.rerun()
    
    st.divider()
    st.header("Your Documents")
    
    documents = vector_store.get_all_documents()
    
    if not documents:
        st.info("No documents uploaded yet. Upload your first document above!")
    else:
        doc_ids = [doc['id'] for doc in documents]
        st.checkbox(
            "Select All / Deselect All", 
            key="select_all_docs", 
            on_change=toggle_all_docs, 
            args=(doc_ids,)
        )

        with st.form("delete_documents_form"):
            for doc in documents:
                with st.expander(f"📄 {doc['filename']} ({doc['category']})"):
                    col1, col2, col3, col4 = st.columns([3, 2, 3, 1])
                    with col1:
                        st.write(f"**Type:** {doc['file_type']}")
                    with col2:
                        st.write(f"**Chunks:** {doc['total_chunks']}")
                    with col3:
                        st.write(f"**Uploaded:** {doc['upload_date'].strftime('%Y-%m-%d %H:%M')}")
                    with col4:
                        st.checkbox("Select", key=f"doc_select_{doc['id']}")

            submitted = st.form_submit_button("Delete Selected Documents")
            if submitted:
                ids_to_delete = [doc['id'] for doc in documents if st.session_state.get(f"doc_select_{doc['id']}")]
                
                if not ids_to_delete:
                    st.warning("No documents selected for deletion.")
                else:
                    deleted_count = 0
                    for doc_id in ids_to_delete:
                        try:
                            vector_store.delete_document(doc_id)
                            deleted_count += 1
                        except Exception as e:
                            st.error(f"Failed to delete document with ID {doc_id}: {e}")
                    
                    if deleted_count > 0:
                        st.success(f"Successfully deleted {deleted_count} document(s).")
                    
                    st.rerun()

with tab2:
    st.header("Ask Questions")
    
    documents = vector_store.get_all_documents()
    
    if not documents:
        st.warning("⚠️ No documents in your knowledge base yet. Please upload some documents first!")
    else:
        st.info(f"📚 Knowledge base contains {len(documents)} document(s)")
        
        if 'messages' not in st.session_state:
            st.session_state.messages = []
        
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                if "sources" in message:
                    with st.expander("📑 Sources"):
                        for source in message["sources"]:
                            st.write(f"- {source['filename']} (similarity: {source['similarity']:.2%})")
        
        if prompt := st.chat_input("Ask a question about your documents..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # --- Easter Egg Logic ---
            if prompt.strip() == "?":
                response = "James Leroy Pearson Jr. Foward we go as the flames grow!!"
                with st.chat_message("assistant"):
                    st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response})
            # --- End Easter Egg ---
            else:
                with st.chat_message("assistant"):
                    with st.spinner("Searching knowledge base..."):
                        query_embedding = embedding_service.generate_query_embedding(prompt)
                        
                        context_chunks = vector_store.search_similar(
                            query_embedding, 
                            top_k=config.TOP_K_RESULTS
                        )
                        
                        if not context_chunks:
                            response = "I couldn't find any relevant information in your knowledge base to answer this question."
                        else:
                            response = generator_service.generate_response(prompt, context_chunks)
                        
                        st.markdown(response)
                        
                        if context_chunks:
                            with st.expander("📑 Sources"):
                                for chunk in context_chunks:
                                    st.write(f"- {chunk['filename']} (similarity: {chunk['similarity']:.2%})")
                
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": response,
                    "sources": [{"filename": c["filename"], "similarity": c["similarity"]} for c in context_chunks] if context_chunks else []
                })

with tab3:
    st.header("Automated Document Ingestion")

    input_path = st.text_input("Input Folder Path", key="input_path_auto")
    output_path = st.text_input("Output Folder Path", key="output_path_auto")

    if st.button("Start Processing", type="primary", key="auto_ingestion_process"):
        if not input_path or not output_path:
            st.error("Please provide both input and output folder paths.")
        elif not os.path.isdir(input_path):
            st.error(f"Input path is not a valid directory: {input_path}")
        elif not os.path.isdir(output_path):
            st.error(f"Output path is not a valid directory: {output_path}")
        else:
            st.info(f"Starting processing from '{input_path}'...")
            
            allowed_types = ['pdf', 'docx', 'txt']
            files_to_process = []
            for ext in allowed_types:
                files_to_process.extend(glob.glob(os.path.join(input_path, f"*.{ext}")))

            if not files_to_process:
                st.warning("No supported files found in the input directory.")
            else:
                total_files = len(files_to_process)
                progress_bar = st.progress(0, text=f"Found {total_files} files to process.")

                for i, file_path in enumerate(files_to_process):
                    file_name = os.path.basename(file_path)
                    progress_text = f"Processing file {i + 1}/{total_files}: {file_name}"
                    progress_bar.progress((i) / total_files, text=progress_text)

                    try:
                        file_type = file_name.split('.')[-1].lower()
                        with open(file_path, "rb") as f:
                            text = document_processor.extract_text(f, file_type)
                        
                        chunks = [text[j:j+2000] for j in range(0, len(text), 2000)]
                        
                        for chunk_index, chunk in enumerate(chunks):
                            doc = Document()
                            header = f"Source: {file_name}"
                            doc.add_paragraph(header)
                            doc.add_paragraph(chunk)
                            
                            output_filename_base = os.path.splitext(file_name)[0]
                            output_filename = f"{output_filename_base}{chunk_index+1}.docx"
                            output_file_path = os.path.join(output_path, output_filename)
                            
                            doc.save(output_file_path)

                        st.success(f"✅ Successfully processed and saved '{file_name}'")

                    except Exception as e:
                        st.error(f"Error processing '{file_name}': {str(e)}")

                progress_bar.progress(1.0, text="All files processed!")
                st.info(f"All processed files have been saved to '{output_path}'")
                st.rerun()

    st.divider()
    st.header("Manage Output Files")

    if not output_path or not os.path.isdir(output_path):
        st.info("Enter a valid output path above to see and manage processed files.")
    else:
        processed_files = glob.glob(os.path.join(output_path, "*.docx"))
        
        if not processed_files:
            st.info("No processed files found in the output directory.")
        else:
            st.checkbox(
                "Select All / Deselect All", 
                key="select_all_output",
                on_change=toggle_all_output_files,
                args=(processed_files,)
            )
            
            with st.form("delete_output_form"):
                for file_path in processed_files:
                    st.checkbox(os.path.basename(file_path), key=f"del_{file_path}")
                    
                submitted = st.form_submit_button("Delete Selected Files")
                if submitted:
                    files_to_delete = [f for f in processed_files if st.session_state.get(f"del_{f}")]

                    if not files_to_delete:
                        st.warning("No files selected for deletion.")
                    else:
                        deleted_count = 0
                        for f_path in files_to_delete:
                            try:
                                os.remove(f_path)
                                deleted_count += 1
                            except Exception as e:
                                st.error(f"Failed to delete {os.path.basename(f_path)}: {e}")
                        
                        if deleted_count > 0:
                            st.success(f"Successfully deleted {deleted_count} file(s).")
                        
                        st.rerun()
